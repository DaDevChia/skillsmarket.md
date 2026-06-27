from __future__ import annotations

import io
from pathlib import Path


class UnsupportedResumeFormat(Exception):
    """Raised when an uploaded file extension is not handled."""


class ResumeExtractionError(Exception):
    """Raised when a supported format could not be parsed into text."""


SUPPORTED_EXTENSIONS = (".txt", ".text", ".pdf", ".docx")


def extract_resume_text(filename: str, raw: bytes) -> str:
    """Extract plain text from an uploaded resume by file extension.

    TXT/PDF/DOCX are supported. Unknown extensions raise
    ``UnsupportedResumeFormat`` (-> 415); parse failures raise
    ``ResumeExtractionError`` (-> 422). Nothing is persisted.
    """
    ext = Path(filename or "").suffix.lower()
    if ext in ("", ".txt", ".text"):
        return raw.decode("utf-8", errors="ignore").strip()
    if ext == ".pdf":
        return _extract_pdf(raw)
    if ext == ".docx":
        return _extract_docx(raw)
    raise UnsupportedResumeFormat(ext or "unknown")


def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        parts = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pypdf raises a variety of errors on bad input
        raise ResumeExtractionError(f"Could not read PDF: {exc}") from exc
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeExtractionError("No selectable text found in PDF (it may be a scan).")
    return text


def _extract_docx(raw: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(raw))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    except Exception as exc:
        raise ResumeExtractionError(f"Could not read DOCX: {exc}") from exc
    text = "\n".join(part for part in parts if part).strip()
    if not text:
        raise ResumeExtractionError("No text found in DOCX.")
    return text
